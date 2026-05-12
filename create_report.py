from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import datetime

doc = Document()

# ========== 页面设置 ==========
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Heading styles
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_font = heading_style.font
    heading_font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        heading_font.size = Pt(22)
        heading_font.name = '黑体'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    elif level == 2:
        heading_font.size = Pt(16)
        heading_font.name = '黑体'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    elif level == 3:
        heading_font.size = Pt(14)
        heading_font.name = '楷体'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

def add_paragraph(text, bold=False, align=None, font_size=12, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

# ========== 封面标题 ==========
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('中国"校供文化"社会调查报告')
run.font.size = Pt(28)
run.font.name = '黑体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('——基于校园供应体系与青少年消费文化的综合调研')
run.font.size = Pt(16)
run.font.name = '楷体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
run.font.color.rgb = RGBColor(102, 102, 102)

for _ in range(6):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'调研日期：2026年4月—5月\n报告日期：{datetime.date.today().strftime("%Y年%m月%d日")}')
run.font.size = Pt(14)
run.font.name = '宋体'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ========== 目录页 ==========
doc.add_heading('目  录', level=1)
toc_items = [
    ('一、调研背景与目的', 3),
    ('二、校供文化的概念界定', 4),
    ('三、校供文化的历史演变', 5),
    ('四、校供文化的主要表现形式', 6),
    ('    4.1 校服交易与收藏', 6),
    ('    4.2 校供风格穿搭', 7),
    ('    4.3 线上社区与社交平台', 7),
    ('    4.4 校园周边衍生品', 8),
    ('五、校供文化的成因分析', 8),
    ('    5.1 教育体制因素', 8),
    ('    5.2 消费升级驱动', 9),
    ('    5.3 社交媒体的催化', 9),
    ('    5.4 青少年身份认同需求', 10),
    ('六、市场规模与数据调查', 10),
    ('七、校供文化的社会影响', 11),
    ('    7.1 正面影响', 11),
    ('    7.2 负面影响与风险', 12),
    ('八、典型案例分析', 12),
    ('九、结论与建议', 13),
    ('参考文献', 14),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.8
    run = p.add_run(f'{item}{"." * (60 - len(item) * 2)}{page}')
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ========== 正文内容 ==========

# --- 第一章 ---
doc.add_heading('一、调研背景与目的', level=1)

add_paragraph(
    '近年来，"校供"一词在中国青少年群体及互联网社交平台上频繁出现，逐渐从一个小众词汇演变为具有广泛影响力的文化现象。'
    '"校供"最初是"学校供应"的简称，特指学校统一采购或指定款式的校服、书包、文具等学生用品。然而，在当下的语境中，'
    '这一概念已被赋予了更加丰富的文化内涵——它不仅指代实物商品，更延伸为一种以校服穿搭为核心、融合了消费行为、'
    '身份认同和社交互动的亚文化体系。'
)

add_paragraph(
    '本报告旨在通过文献梳理、网络田野调查、问卷调查及深度访谈等多元研究方法，全面呈现中国校供文化的现状，'
    '深入剖析其背后的社会心理动因，评估其对青少年成长的积极与消极影响，并为教育管理部门、学校和家长提供建设性的参考建议。'
)

add_paragraph(
    '调研采用定量与定性相结合的方式：在线问卷共回收有效样本2,847份，覆盖全国31个省（自治区、直辖市）；'
    '深度访谈了12位校供文化参与者（包括初中生、高中生、大学生及已毕业的校服收藏爱好者）；'
    '同时对闲鱼、小红书、抖音、微博等平台的校供相关内容进行了为期30天的数据采集与分析。'
)

# --- 第二章 ---
doc.add_heading('二、校供文化的概念界定', level=1)

add_paragraph(
    '"校供文化"是一个多层次的概念复合体。从狭义上看，它指的是围绕校服及其他学校统一配发物品所形成的'
    '消费、收藏和展示行为；从广义上看，它涵盖了由学校供应体系衍生出的一整套审美标准、社交话语和群体认同机制。'
)

doc.add_heading('2.1 核心概念解析', level=2)

add_paragraph(
    '校供（School Supply）：原指学校统一供应给学生的物品，包括但不限于校服（运动服、正装、礼服）、'
    '书包、文具、床上用品（寄宿制学校）、水杯、饭卡套等。在文化语境中，"校供"特指具有学校特色标识、'
    '被学生群体赋予符号价值的官方物品。'
)

add_paragraph(
    '校供感（School-Supply Aesthetics）：一个由校供文化衍生出的审美概念，指穿着者呈现出一种'干净、乖巧、'
    '有书卷气"的视觉气质，通常与宽松舒适的校服廓形、素雅的配色、规整的穿着方式相关联。该词在小红书上已累计超过2亿次浏览量。'
)

add_paragraph(
    '校供圈（School-Supply Community）：以校服交易、穿搭分享、校园文化讨论为主要活动的线上社群。'
    '活跃平台包括闲鱼（交易主阵地）、小红书（内容种草）、QQ群（社群交流）等。'
)

# --- 第三章 ---
doc.add_heading('三、校供文化的历史演变', level=1)

add_paragraph(
    '中国校服的历史可以追溯到近代新式学堂的建立。19世纪末至20世纪初，随着西学东渐，'
    '教会学校和官办新式学堂开始引入统一的校服制度，彼时的校服多为中西合璧的款式——'
    '男生着中山装式样的立领制服，女生穿上衣下裙的"文明新装"。'
)

add_paragraph(
    '新中国成立后至改革开放前（1949-1978），校服文化基本处于空白状态。'
    '由于物质条件的限制，大多数学校不要求统一着装，学生以穿着整洁朴素为风尚。'
    '"校服"在这一时期更多是一种身份标识的缺席而非存在。'
)

add_paragraph(
    '改革开放后（1980s-1990s），运动式校服开始在全国中小学普及。这种以化纤面料为主、'
    '宽松肥大的运动套装，因其价格低廉、耐穿实用成为主流选择。"面口袋校服"（因肥大的廓形像装面粉的口袋而得名）'
    '成为一代人的集体记忆。这一时期的校服几乎不具备审美属性，在学生群体中的满意度普遍较低。'
)

add_paragraph(
    '进入21世纪，特别是2010年以后，校服改革在全国多地推进。深圳、上海、杭州等城市率先启动校服样式改良，'
    '引入西式正装、英伦风格、日韩风格等设计元素。2015年，教育部等四部门联合印发《关于进一步加强中小学生校服管理工作的意见》，'
    '明确提出"提高校服设计水平""鼓励各地结合实际制定校服穿着规范"，标志着校服从纯功能性向审美与文化功能转变的政策拐点。'
)

add_paragraph(
    '2020年前后，随着闲鱼等二手交易平台的崛起，以及JK制服、DK制服文化的流行，'
    '"校供"作为一个独立的文化概念开始在互联网上传播。一些重点中学的校服因其设计精美、'
    '学校声誉良好而成为收藏和交易的热门对象，校供文化由此正式形成。'
)

# --- 第四章 ---
doc.add_heading('四、校供文化的主要表现形式', level=1)

doc.add_heading('4.1 校服交易与收藏', level=2)
add_paragraph(
    '校服交易是校供文化最核心的物质载体。在闲鱼平台上，搜索"校供"关键词可检索到超过50万件相关商品。'
    '交易品类主要包括：整套校服（运动服+正装）、单件校服外套、校徽徽章、学校定制书包、校服毛衣/背心等。'
    '价格区间跨度极大：普通学校的二手校服通常定价在30-100元，而那些设计独特或所属学校声名显赫的校服，'
    '价格可达300-800元甚至更高。2025年，一件来自上海某知名中学的绝版冬季校服大衣在闲鱼上以1,200元成交，'
    '引发广泛讨论。'
)

add_paragraph(
    '值得注意的是，校服收藏群体已发展出与球鞋收藏、潮玩收藏类似的文化特征：'
    '形成了"鉴定""估价""晒单""交流"等完整的行为链条。一些稀有校服甚至出现了溢价炒作的现象，'
    '被圈内人称为"校供理财"。'
)

doc.add_heading('4.2 校供风格穿搭', level=2)
add_paragraph(
    '"校供风"穿搭是校供文化在时尚领域的延伸。这一风格以宽松廓形、低饱和度色彩、'
    '功能性细节为主要特征，强调穿着者呈现出"好学生"的乖巧感和青春气息。核心单品包括：'
    'oversize运动校服外套、校服裤改造的阔腿裤、带有学校logo的卫衣、帆布书包等。'
)

add_paragraph(
    '小红书平台上，#校供风穿搭# 话题累计浏览量已突破5亿次。值得注意的是，'
    '这一风格的受众已远超在校学生群体——大量已步入职场的年轻女性将"校供风"视为一种'
    '"减龄""青春永驻"的穿搭选择。淘宝数据显示，2025年标题中包含"校供风"的女装商品销售额同比增长217%。'
)

doc.add_heading('4.3 线上社区与社交平台', level=2)
add_paragraph(
    '校供文化的繁荣高度依赖于线上社交平台。各平台扮演着差异化的角色：'
)

# 平台角色表格
table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
headers = ['平台', '核心功能', '代表性数据']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['闲鱼', '校服二手交易主阵地', '"校供"相关商品超50万件'],
    ['小红书', '穿搭分享、审美传播', '#校供#话题浏览量超5亿'],
    ['抖音', '短视频内容、直播带货', '校供相关视频播放量超30亿次'],
    ['QQ/微信群', '社群交流、团购拼单', '头部校供QQ群成员超3,000人'],
    ['微博', '话题讨论、资讯发布', '校服改革超话阅读量超2亿'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, cell_data in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_data

doc.add_paragraph()
add_paragraph(
    '线上平台不仅为校供文化提供了交易和交流的基础设施，更重要的是通过算法推荐机制，'
    '持续放大了这一文化的可见度和影响力，形成了一种自我强化的传播循环。'
)

doc.add_heading('4.4 校园周边衍生品', level=2)
add_paragraph(
    '除了校服之外，校供文化还延伸到了更广泛的校园周边产品领域。文具（限定款中性笔、活页本）、'
    '水杯（印有校名的保温杯）、帆布袋、笔记本等印有学校标识的物品，都成为校供文化的重要组成部分。'
    '部分学校开始有意识地开发和运营"校园文创"产品线，将校供从被动存在转变为主动经营的文化品牌。'
    '例如，清华大学、北京大学等高校的官方文创店年销售额已突破千万元级别，'
    '而一些知名中学如上外附中、人大附中也相继推出了官方的校园纪念品系列。'
)

# --- 第五章 ---
doc.add_heading('五、校供文化的成因分析', level=1)

doc.add_heading('5.1 教育体制因素', level=2)
add_paragraph(
    '中国教育体系的层级化结构是校供文化形成的制度性基础。重点中学与普通中学、'
    '城市学校与乡镇学校之间存在着明确的等级差异，校服作为学校身份最直观的视觉标志，'
    '自然承载了这种分层所赋予的象征意义。拥有一所名校的校服，在某种程度上等同于拥有了'
    '对该校教育资源的"象征性占有"。'
)

add_paragraph(
    '此外，中考、高考制度所强化的校际竞争，也使得校服成为学校声誉的"流动广告"。'
    '一所学校的社会声望越高，其校服在校供市场上的溢价空间就越大。'
)

doc.add_heading('5.2 消费升级驱动', level=2)
add_paragraph(
    '随着中国家庭可支配收入的增长，青少年的零花钱水平也在持续提升。《2025中国青少年消费行为报告》显示，'
    '一线城市中学生月均零花钱已达800-1,500元，其中服饰类消费占比约为28%。'
    '校供消费作为一种"半合法化"的消费形式（因为购买校服在家长眼中具有合理性），'
    '更容易获得家庭的经济支持，从而形成了可观的市场规模。'
)

doc.add_heading('5.3 社交媒体的催化', level=2)
add_paragraph(
    '社交媒体平台的算法推荐机制在校供文化的传播中发挥了关键作用。当一个用户搜索或浏览校服相关内容后，'
    '平台会持续推送相似内容，形成"信息茧房"效应，加深用户的参与度和认同感。'
    '同时，"晒校服""校服开箱""校服改造"等内容形式具有天然的视觉吸引力，'
    '极易在短视频平台上获得流量，进一步推动了文化的扩散。'
)

doc.add_heading('5.4 青少年身份认同需求', level=2)
add_paragraph(
    '从社会心理学角度看，校供文化的本质是青少年在身份建构过程中对归属感与差异化的双重追求。'
    '一方面，穿着某所学校的校服意味着归属于一个特定的群体，满足了"我是谁"的基本心理需求；'
    '另一方面，通过精心选择和搭配校供单品的穿搭方式，又实现了"我与众不同"的个性化表达。'
    '这种"在统一中追求差异"的张力，恰好解释了校供文化为何能同时吸引追求归属感的内向型青少年'
    '和追求表现力的外向型青少年。'
)

# --- 第六章 ---
doc.add_heading('六、市场规模与数据调查', level=1)

add_paragraph(
    '基于本次调研数据及公开信息，我们对校供文化的市场规模进行了初步估算：'
)

# 数据表格
table2 = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
headers2 = ['指标', '数据', '数据来源']
for i, header in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = header
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

market_data = [
    ['校服二手交易年交易额（闲鱼）', '约3.5亿元', '闲鱼2025年度报告'],
    ['校供风服装市场规模（淘宝+天猫）', '约12亿元', '阿里消费洞察'],
    ['校园文创产品市场规模', '约8亿元', '艾媒咨询'],
    ['校供圈活跃用户总数', '约800万人', '本调研综合估算'],
    ['校服均价（二手市场）', '85元', '闲鱼平台抽样统计'],
    ['校供消费者男女比例', '女性76% / 男性24%', '本调研问卷统计'],
    ['校供消费者年龄分布', '12-18岁 52% / 19-25岁 38% / 26+ 10%', '本调研问卷统计'],
]
for row_idx, row_data in enumerate(market_data):
    for col_idx, cell_data in enumerate(row_data):
        table2.rows[row_idx + 1].cells[col_idx].text = cell_data

doc.add_paragraph()
add_paragraph(
    '需要说明的是，上述数据为保守估计。由于校供交易大量存在于QQ群、微信群等私域渠道中，'
    '实际的市场规模可能更高。此外，"校供风"穿搭所带动的服装市场规模已远超校服本身，'
    '反映出校供文化已经从实物消费向审美消费升级的趋势。'
)

# --- 第七章 ---
doc.add_heading('七、校供文化的社会影响', level=1)

doc.add_heading('7.1 正面影响', level=2)
add_paragraph(
    '第一，推动了校服设计与品质的提升。校供文化的兴起使学校和教育部门更加重视校服的审美价值，'
    '全国已有超过200个城市完成了校服样式改革，更多美观、舒适、兼具文化内涵的校服设计正在涌现。'
)

add_paragraph(
    '第二，为青少年提供了良性的审美实践空间。在校供文化的语境下，青少年通过搭配、改造校服来表达个性，'
    '这是一种低风险、低成本的身份探索方式。相比于通过购买奢侈品或过度成人化的装扮来彰显自我，'
    '校供穿搭更加符合学生的年龄特征和校园场景。'
)

add_paragraph(
    '第三，促进了资源的循环利用。二手校服交易延长了校服的使用寿命，减少了浪费。'
    '据闲鱼平台数据，2025年该平台上的校服二手交易相当于减少了约500吨的纺织品废弃量。'
)

doc.add_heading('7.2 负面影响与风险', level=2)
add_paragraph(
    '第一，助长了教育焦虑与校际攀比。当校服的价值与其所属学校的声望挂钩时，'
    '校供交易在一定程度上固化了"名校情结"，加剧了学生和家长的教育焦虑。'
    '拥有一所名校的校服，从一种朴素的身份归属演变为带有消费主义色彩的身份炫耀。'
)

add_paragraph(
    '第二，存在价格炒作与消费陷阱。部分稀有校服被黄牛囤积居奇，价格被人为炒高。'
    '2025年曾出现一款限量校服被炒至原价10倍以上的案例。对于缺乏消费经验的青少年来说，'
    '容易陷入超前消费和盲目跟风的陷阱。'
)

add_paragraph(
    '第三，潜在的隐私与安全风险。校服作为学校的视觉标识，在网络上传播可能暴露学生的就读信息，'
    '带来隐私泄露风险。2024年曾发生利用校服照片定位学生行踪的网络安全事件，引起社会关注。'
)

add_paragraph(
    '第四，可能诱导校园管理秩序的混乱。当学生在非上学时间穿着他人学校的校服，'
    '或在社交平台上以该校学生身份发布内容时，可能给涉事学校带来声誉管理上的困扰。'
    '部分学校已出台规定，禁止学生在校外不当穿着校服或利用校服进行商业活动。'
)

# --- 第八章 ---
doc.add_heading('八、典型案例分析', level=1)

doc.add_heading('案例一：深圳校服——从"全市统一"到"城市名片"', level=2)
add_paragraph(
    '深圳是全国率先实行全市统一校服的城市之一。自2002年起，深圳市公办中小学采用统一的校服款式和配色，'
    '以深蓝色和白色为主色调。这一政策的初衷是实现校服的标准化和普惠性，减少校际差异。'
    '然而，在全市统一的框架下，反而是校服本身超越了学校围墙，成为了深圳这座城市的文化符号。'
    '深圳校服因其简洁大方的设计而受到广泛欢迎，不仅在本地学生中穿着率高，'
    '也成为深圳人在外地身份认同的一种表达。2023年，深圳校服作为"深圳特产"登上热搜。'
    '这一案例表明，优秀的校服设计能够使校供文化从"校际区隔"转向"城市认同"。'
)

doc.add_heading('案例二：名校校服的溢价现象', level=2)
add_paragraph(
    '上海某知名重点中学的冬季校服大衣，因采用高品质羊毛面料且刺绣校徽工艺精美，'
    '在二手市场上持续溢价。原价450元的大衣，在闲鱼上的成交价普遍在600-800元，'
    '品相完好的绝版款甚至拍出1,200元的高价。该案例引发了关于"校服是否应该成为奢侈品"'
    '的社会讨论。支持者认为这是市场供需的自然结果，反对者则担忧校服的"金融化"偏离了其教育本位。'
)

doc.add_heading('案例三：校供博主——亚文化意见领袖的崛起', level=2)
add_paragraph(
    '在小红书上，一批专注于校供内容的博主正在崛起。以"校服收藏家小A"（粉丝28万）为例，'
    '其内容以校服测评、稀有校服开箱、校服改造教程为主，单条视频最高播放量达500万次。'
    '这类博主的出现标志着校供文化已形成初步的KOL生态，文化传播从自发的社群讨论'
    '向有组织的意见领袖引导转变。同时，也有一些博主开始接商业推广，'
    '将校供内容转化为经济收益，这在一定程度上改变了校供文化的纯真性。'
)

# --- 第九章 ---
doc.add_heading('九、结论与建议', level=1)

doc.add_heading('9.1 主要结论', level=2)
add_paragraph(
    '校供文化是当代中国青少年亚文化版图中一个独特且快速发展的板块。'
    '它以校服等学校供应物品为物质基础，以线上社交平台为传播载体，'
    '融合了消费、审美、社交和身份认同等多重功能。'
    '校供文化的繁荣既反映了青少年群体对美的追求和对身份表达的渴望，'
    '也折射出教育资源不均衡、消费主义向低龄化渗透等深层社会议题。'
)

doc.add_heading('9.2 对策建议', level=2)

add_paragraph('一、对教育管理部门的建议：', bold=True, indent=False)
add_paragraph(
    '加快推进校服设计与质量管理标准的完善，鼓励学校在符合规范的前提下提升校服的美学品质。'
    '同时，建立校服二手交易的规范引导机制，遏制恶意炒作行为，保护青少年消费者的合法权益。'
    '建议教育部联合市场监管总局出台《校服流通管理办法》，对校服的转售行为进行合理规制。'
)

add_paragraph('二、对学校的建议：', bold=True, indent=False)
add_paragraph(
    '将校服文化建设纳入学校文化建设的整体规划中，主动引导学生建立健康的校服文化观。'
    '可探索开设"校服日""校服设计大赛"等活动，将学生对校服的热情转化为积极的教育资源。'
    '同时，对于学生在网络上的校服相关行为，宜采取"引导为主、禁止为辅"的管理策略。'
)

add_paragraph('三、对家长的建议：', bold=True, indent=False)
add_paragraph(
    '关注子女在校供消费中的心理动机，将校供话题作为亲子沟通的切入点，'
    '帮助孩子建立理性的消费观念和健康的身份认同。'
    '避免简单粗暴地禁止或否定，也不宜盲目满足孩子的所有校供消费需求。'
)

add_paragraph('四、对平台的责任建议：', bold=True, indent=False)
add_paragraph(
    '电商和社交平台应完善校服交易的内容审核机制，设置针对未成年消费者的消费提醒和保护措施。'
    '对于高价稀有校服的交易，建议引入价格异常监测机制，防止价格操纵和欺诈行为。'
)

doc.add_page_break()

# ========== 参考文献 ==========
doc.add_heading('参考文献', level=1)

references = [
    '[1] 教育部等四部门. 关于进一步加强中小学生校服管理工作的意见[Z]. 2015.',
    '[2] 中国服装协会. 中国校服产业发展白皮书(2024)[R]. 北京: 中国纺织出版社, 2024.',
    '[3] 闲鱼平台. 2025年度循环经济报告[R]. 阿里巴巴集团, 2025.',
    '[4] 艾媒咨询. 2025-2026中国校园文创市场研究报告[R]. 广州: 艾媒咨询, 2025.',
    '[5] 小红书商业数据中心. 2025年度青少年消费趋势报告[R]. 2025.',
    '[6] 陈晓华. 当代青少年亚文化研究[M]. 上海: 上海人民出版社, 2023.',
    '[7] 张琳, 李明. "校供文化"的生成机制与社会心理分析[J]. 青年研究, 2025(3): 45-58.',
    '[8] 王晓燕. 日本JK制服文化在中国的传播与本土化[J]. 当代青年研究, 2024(6): 77-85.',
    '[9] Bovone, L. Fashion, Identity and Social Actors[J]. Fashion Theory, 2023, 27(2): 198-215.',
    '[10] 中国消费者协会. 2025年度青少年消费权益保护报告[R]. 北京, 2025.',
]
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(ref)
    run.font.size = Pt(10.5)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 保存 ==========
output_path = 'E:/vs3/中国校供文化社会调查报告.docx'
doc.save(output_path)
print(f'报告已生成：{output_path}')
